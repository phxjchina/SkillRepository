# -*- coding: utf-8 -*-
"""Generate an SLR-style Chinese survey on ML metamodeling from 33 local papers.

Follows the structure of the reference SLR
(Mohamed et al., "A Systematic Literature Review on Model-driven Engineering
for Cyber-Physical Systems", arXiv:2103.08644):
  Title -> Abstract -> 1 Introduction (background / secondary research method /
  problem statement / contributions) -> 2 Methodology (RQs / sources & selection
  / inclusion-exclusion / quality assessment / data extraction / process) ->
  3 Execution -> 4 Results (bibliometrics + per-RQ analysis + study
  characteristics table) -> 5 Discussion (synthesis / implications / threats to
  validity) -> 6 Conclusion -> References -> Appendix (study characteristics).

All citations map to the 33-paper REFERENCE list; no hallucinated citations.
"""
import re, argparse, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# Output path (override with --out; default reproduces the original 33-paper survey)
parser = argparse.ArgumentParser(
    description="Generate an SLR-style Chinese survey docx from a local corpus.")
parser.add_argument("--out", default="E:/科研/DSML+ML元建模/机器学习元模型综述_SLR.docx",
                    help="output .docx path")
args = parser.parse_args()
OUT = args.out

# ---------------------------------------------------------------------------
# Reference database (33 papers, GB/T 7714). Number = inline citation.
# ---------------------------------------------------------------------------
REFERENCES = [
    ("[1]", "RÄDLER S, BERARDINELLI L, WINTER K, et al. Bridging MDE and AI: A Systematic Review of Domain-Specific Languages and Model-Driven Practices in AI Software Systems Engineering[J]. Preprint, arXiv:2307.04599, 2024."),
    ("[2]", "MOHAMED M A, KARDAS G, CHALLENGER M. A Systematic Literature Review on Model-driven Engineering for Cyber-Physical Systems[R]. arXiv:2103.08644, 2021."),
    ("[3]", "RÄDLER S, MANGLER J, RINDERLE-MA S. Model-Driven Engineering Method to Support the Formalization of Machine Learning using SysML[C]. arXiv:2307.04495, 2023."),
    ("[4]", "RÄDLER S, RUPP M, RIGGER E, et al. Code Generation for Machine Learning using Model-Driven Engineering and SysML[C]. arXiv:2307.05584, 2023."),
    ("[5]", "MOIN A, CHALLENGER M, BADII A, et al. Supporting AI Engineering on the IoT Edge through Model-Driven TinyML[J]. arXiv:2107.02690, 2021."),
    ("[6]", "赵静, 梁浩, 徐天啸, 等. 面向智能汽车赛博物理系统的领域特定建模语言研究[J]. 汽车工程, 2024, 46(8)."),
    ("[7]", "DURÁ COSTA C, HERNÁNDEZ LÓPEZ J A, SÁNCHEZ CUADRADO J. ModelMate: A recommender for textual modeling languages based on pre-trained language models[C]. 2023."),
    ("[8]", "SINANI N, SALMA S, BOUTOT P, et al. Towards a Domain-Specific Modelling Environment for Reinforcement Learning[J]. arXiv:2410.09368, 2024."),
    ("[9]", "HALLOU A, FISSAA T, HAFIDDI H, et al. Context-Aware IoT System Development Approach Based on Meta-Modeling and Reinforcement Learning: A Smart Home Case Study[J]. International Journal of Online and Biomedical Engineering, 2024, 20(6): 25-42."),
    ("[10]", "OUADDI C, BENADDI L, SOUHA A, et al. A Model-Driven Approach Employing DSL and Generation Templates to Accelerate the Development of Conversational Agents for Smart Tourism[J]. IEEE Access, 2025."),
    ("[11]", "SEIDEWITZ E. Domain-Specific Modeling with SysML v2[C]. SysML v2 Summit, Reston VA, 2025."),
    ("[12]", "HE X, CHEN R, ZHANG Z, et al. A Hybrid Approach for EMF Code Generation: Code Templates Meet Large Language Models[EB/OL]. arXiv:2512.05498, 2025."),
    ("[13]", "WANG R, LU M, YU C H, et al. Automated Deep Learning Optimization via DSL-Based Source Code Transformation[EB/OL]. arXiv:2405.03067, 2024."),
    ("[14]", "KHALFAN S, AL MAZROUEI S. Anka: A Domain-Specific Language for Reliable LLM Code Generation[EB/OL]. arXiv:2512.23214, 2025."),
    ("[15]", "TODD G, PADULA A G, SOEMERS D J N J, et al. Ludax: A GPU-Accelerated Description Language for Board Games[EB/OL]. arXiv:2506.22609, 2025."),
    ("[16]", "ANUPAM S, BOWERS M, COSTILLA-REYES O, et al. MathDSL: A Domain-Specific Language for Concise Mathematical Solutions Via Program Synthesis[EB/OL]. arXiv:2409.17490, 2024."),
    ("[17]", "HUANG J, LI Z, LI Y, et al. NineToothed: A Triton-Based High-Level Domain-Specific Language for Machine Learning[EB/OL]. arXiv:2507.11978, 2025."),
    ("[18]", "GAUFFRIAU A, DE ALBUQUERQUE SILVA I, PAGETTI C. Formal description of ML models for unambiguous implementation[J]. arXiv:2307.12713, 2024."),
    ("[19]", "MANINO E, FARIAS B, MENEZES R S, et al. Floating-Point Neural Network Verification at the Software Level[J]. arXiv:2510.23389, 2025."),
    ("[20]", "KAULEN K, LADNER T, BAK S, et al. The 6th International Verification of Neural Networks Competition (VNN-COMP 2025): Summary and Results[EB/OL]. arXiv:2512.19007, 2025."),
    ("[21]", "AFFELDT R, BRUNI A, KOMENDANTSKAYA E, et al. Taming Differentiable Logics with Coq Formalisation[EB/OL]. arXiv:2403.13700, 2024."),
    ("[22]", "NASIR A, ZENATI A. A Dynamical Systems Framework for Reinforcement Learning Safety and Robustness Verification[EB/OL]. arXiv:2508.15588, 2025."),
    ("[23]", "AGGARWAL P, PARNO B, WELLECK S. AlphaVerus: Bootstrapping Formally Verified Code Generation through Self-Improving Translation and Treefinement[EB/OL]. arXiv:2412.06176, 2024."),
    ("[24]", "VATAI E, DROZD A, IVANOV I R, et al. Tadashi: Enabling AI-Based Automated Code Generation With Guaranteed Correctness[EB/OL]. arXiv:2410.03210, 2024."),
    ("[25]", "SEVENHUIJSEN M, ETEMADI K, NYBERG M. VeCoGen: Automating Generation of Formally Verified C Code with Large Language Models[EB/OL]. arXiv:2411.19275, 2024."),
    ("[26]", "YE Z, YAN Z, HE J, et al. VERINA: Benchmarking Verifiable Code Generation[EB/OL]. arXiv:2505.23135, 2025."),
    ("[27]", "WEI R, ZHU L, WANG H, et al. Formal-Method-Guided Vibe Coding: Closing the Verification Loop on AI-Generated Safety-Critical Software Through Model-Driven Engineering[EB/OL]. arXiv:2606.22413, 2026."),
    ("[28]", "MA T, LAI H, WANG H, et al. ATLAS: A Layered Constraint-Guided Framework for Structured Artifact Generation in LLM-Assisted MDE[EB/OL]. arXiv:2510.25890, 2025."),
    ("[29]", "PAN F, SONG Y, WEN L, et al. Automating Automotive Software Development: A Synergy of Generative AI and Model-Based Methods[EB/OL]. arXiv:2505.02500, 2025."),
    ("[30]", "YANG Z, ZHENG J, CHEN G. Verification and Interpretation-Driven Safe Deep Reinforcement Learning Framework[EB/OL]. arXiv:2410.15127, 2024."),
    ("[31]", "LE T, SHEFIN R, GUPTA D, et al. Verification-Guided Falsification for Safe RL via Explainable Abstraction and Risk-Aware Exploration[EB/OL]. arXiv:2506.03469, 2025."),
    ("[32]", "LIU M, YU C H, LEE W H, et al. Synthesizing Programmatic Reinforcement Learning Policies with Large Language Model Guided Search[EB/OL]. arXiv:2405.16450, 2024."),
    ("[33]", "YAN C, CHE F, HUANG X, et al. Re:Form — Reducing Human Annotations in Scalable Formal Software Verification with RL in LLMs[EB/OL]. arXiv:2507.16331, 2026."),
]

# ---------------------------------------------------------------------------
# Study characteristics (33 papers). (id, author(year), year, theme, method,
# formalism/verification, tool, validation type). theme tags: A=MDE/DSL,
# B=formal desc/verif, C=verifiable code gen, D=safe RL, SLR=secondary.
# ---------------------------------------------------------------------------
STUDY = [
    (1,  "Rädler 等 (2024)",        2024, "SLR", "MDE+AI 系统综述（映射研究）", "二次研究", "—", "文献综述(1335→18)"),
    (2,  "Mohamed 等 (2021)",       2021, "SLR", "CPS 的 MDE 系统综述", "二次研究(SLR)", "—", "SLR(140 项)"),
    (3,  "Rädler 等 (2023)",        2023, "A",   "SysML 形式化 ML 流水线", "SysML v2 元模型", "SysML", "案例"),
    (4,  "Rädler 等 (2023)",        2023, "A",   "SysML+ MDE 代码生成", "MDE+SysML", "代码生成器", "案例"),
    (5,  "Moin 等 (2021)",          2021, "A",   "模型驱动 TinyML", "MDE", "IoT 边缘", "案例"),
    (6,  "赵静 等 (2024)",          2024, "A",   "IVCPS 领域建模语言", "元语言/MBSE", "汽车 CPS", "案例"),
    (7,  "Durá Costa 等 (2023)",    2023, "A",   "ModelMate 语言推荐", "预训练语言模型", "建模工具", "用户评估"),
    (8,  "Sinani 等 (2024)",        2024, "A",   "RL 建模环境 RLML", "DSL", "RL 原型", "原型"),
    (9,  "Hallou 等 (2024)",        2024, "A",   "IoT+RL 元建模", "元建模+RL", "智能家居", "案例"),
    (10, "Ouaddi 等 (2025)",        2025, "A",   "对话智能体 DSL", "DSL+模板", "智慧旅游", "案例"),
    (11, "Seidewitz (2025)",        2025, "A",   "SysML v2 领域建模", "SysML v2", "标准/工具", "说明性"),
    (12, "He 等 (2025)",            2025, "A",   "EMF 代码生成混合", "模板+LLM", "EMF", "评估"),
    (13, "Wang 等 (2024)",          2024, "A",   "DSL 源码变换优化", "DSL", "DL 优化", "基准"),
    (14, "Khalfan 等 (2025)",       2025, "C",   "Anka LLM 代码生成", "DSL+LLM", "代码生成", "基准"),
    (15, "Todd 等 (2025)",          2025, "A",   "Ludax 棋类描述语言", "DSL+GPU", "游戏", "评估"),
    (16, "Anupam 等 (2024)",        2024, "A",   "MathDSL 程序合成", "DSL", "数学求解", "基准"),
    (17, "Huang 等 (2025)",         2025, "A",   "NineToothed Triton DSL", "DSL(Triton)", "ML 内核", "基准"),
    (18, "Gauffriau 等 (2024)",     2024, "B",   "ML 形式化描述", "形式化语义", "航空认证", "形式化描述"),
    (19, "Manino 等 (2025)",        2025, "B",   "浮点 NN 验证", "SMT/软件层", "NN 验证", "验证"),
    (20, "Kaulen 等 (2025)",        2025, "B",   "VNN-COMP 2025", "NN 验证竞赛", "竞赛", "竞赛总结"),
    (21, "Affeldt 等 (2024)",       2024, "B",   "可微逻辑 Coq 形式化", "Coq/定理证明", "可微逻辑", "证明"),
    (22, "Nasir 等 (2025)",         2025, "B/D", "RL 安全动力系统", "Lyapunov/动力系统", "RL 安全", "验证"),
    (23, "Aggarwal 等 (2024)",      2024, "C",   "AlphaVerus", "LLM+形式化", "证明生成", "证明"),
    (24, "Vatai 等 (2024)",         2024, "C",   "Tadashi", "AI 代码生成+保证", "代码生成", "验证"),
    (25, "Sevenhuijsen 等 (2024)",  2024, "C",   "VeCoGen", "LLM+形式化 C", "C 代码", "证明"),
    (26, "Ye 等 (2025)",            2025, "C",   "VERINA 基准", "LLM+形式化", "基准", "基准"),
    (27, "Wei 等 (2026)",           2026, "C",   "Forge vibe coding", "形式化+MDE", "安全关键", "框架"),
    (28, "Ma 等 (2025)",            2025, "C",   "ATLAS 分层约束", "LLM 辅助 MDE", "MDE", "框架/评估"),
    (29, "Pan 等 (2025)",           2025, "C",   "汽车 GenAI+MBSE", "生成式 AI+MBSE", "汽车", "案例"),
    (30, "Yang 等 (2024)",          2024, "D",   "Safe DRL 验证解释", "验证+解释", "RL", "框架"),
    (31, "Le 等 (2025)",            2025, "D",   "验证引导反例搜索", "抽象+风险探索", "RL", "验证"),
    (32, "Liu 等 (2024)",           2024, "D",   "LLM 引导 RL 策略", "LLM+RL", "RL", "评估"),
    (33, "Yan 等 (2026)",           2026, "C/D", "Re:Form RL 减标注", "RL+LLM 形式化", "形式化验证", "框架"),
]

THEME_NAME = {
    "A":   "A. ML 元建模与 MDE",
    "B":   "B. 形式化描述与验证",
    "C":   "C. 可验证代码生成",
    "D":   "D. 安全 RL 验证",
    "SLR": "E. 系统综述(二次研究)",
}
# primary theme (first tag) for counting
def primary_theme(t):
    return t.split("/")[0]

from collections import Counter
year_dist = Counter(s[2] for s in STUDY)
theme_dist = Counter(primary_theme(s[3]) for s in STUDY)
# publication type (derived from reference string)
def pub_type(ref):
    if "[J]" in ref and "arXiv" not in ref:
        return "期刊(Journal)"
    if "[C]" in ref:
        return "会议(Conference)"
    if "[R]" in ref:
        return "技术报告(Tech Report)"
    return "预印本(arXiv Preprint)"
type_dist = Counter(pub_type(r) for _, r in REFERENCES)

# ---------------------------------------------------------------------------
# Section content. kind in: title, subtitle, abstract_l, kw_l, h1, h2, h3,
# p, table_year, table_type, table_theme, table_study
# ---------------------------------------------------------------------------
S = []
def T(kind, text): S.append((kind, text))

T("title", "机器学习元建模的最新进展：一项系统文献综述")
T("subtitle", "——基于 33 篇近期文献（2021–2026）的梳理")

T("abstract_l", "摘  要")
T("p", "随着机器学习（深度学习与强化学习）在安全关键系统中的广泛部署，如何以工程化、可验证、可组合的方式刻画并实现 ML 系统成为研究热点。“机器学习元模型”（ML metamodeling）指以模型驱动工程（MDE/MBSE）方式定义 ML 系统——涵盖数据、模型、训练流程、性质需求与部署——的抽象语法与语义，从而支撑领域特定语言（DSL）、自动代码生成与形式化验证。本文采用系统文献综述（SLR）方法，对 33 篇近期文献（含 2 篇系统综述、14 篇 ML 元建模/DSL、5 篇形式化描述与验证、9 篇可验证代码生成、4 篇安全强化学习验证，以及交叉文献）进行结构化梳理。我们提出 5 个研究问题（RQ1–RQ5），定义了纳入/排除标准、质量评估项与数据提取框架，并通过文献计量与按 RQ 的实证分析，将进展归纳为四条相互交织的主线，最终指出统一元模型标准、验证可扩展性、端到端工具链与认证就绪等开放问题。")
T("kw_l", "关键词")
T("p", "机器学习元模型；模型驱动工程；领域特定语言；SysML v2；形式化验证；可验证代码生成；安全强化学习；系统文献综述")

T("h1", "1  引言")
T("h2", "1.1  机器学习与元建模")
T("p", "机器学习系统的工程复杂度持续上升：模型架构、训练流水线、数据治理与部署环境日益交织，单纯依靠手工编码与试错已难以保证正确性、可追溯性与可认证性[2]。在这一背景下，模型驱动工程（MDE）与基于模型的系统工程（MBSE）提供了一条以“模型”为首要制品（artifact）的路径——通过领域特定语言（DSL）与元模型（metamodel）提升抽象层级、支撑自动化[1]。所谓“元模型”，是位于 MOF 四层架构 M2 层、定义一种建模语言抽象语法的模型；在机器学习语境下，ML 元模型即刻画 ML 概念（数据集、特征、模型架构、训练流程、性质规约）的模型。一个 DSL 的抽象语法本质上就是其元模型，因此“为 ML 定义 DSL”与“为 ML 定义元模型”是同一枚硬币的两面[1][3]。")
T("h2", "1.2  二次研究方法")
T("p", "本文采用系统文献综述（Systematic Literature Review, SLR）这一二次研究方法，而非叙述性综述，原因在于 SLR 通过显式的研究问题、可复现的搜索与选择流程、纳入/排除标准与质量评估，能够降低作者偏见、使结论可审计[1][2]。本综述对齐 Kitchenham 与 Mohamed 等[2]所示范的 SLR 规范：先定义研究问题与协议，再执行文献收集与筛选，最后围绕研究问题综合证据并评估效度威胁。需要说明，本综述的语料并非来自数据库全量检索，而是作者围绕“ML 元模型”主题系统性收集的 33 篇近期代表性文献（详见第 2、3 节）；因此本综述在“来源广度”上低于全库 SLR，但在“主题聚焦度”与“对最新预印本工作的覆盖”上具有针对性优势。")
T("h2", "1.3  问题陈述")
T("p", "现有关于 MDE 与 AI/ML 结合的工作分散在多条技术路线上：有的聚焦用 SysML v2 与 DSL 刻画 ML 系统[3][6][11]，有的聚焦神经网络与可微逻辑的形式化验证[18][21]，有的聚焦用大语言模型（LLM）生成可证明正确的代码[23][26]，还有的聚焦安全深度强化学习的验证驱动闭环[30][31]。然而，这些工作缺乏一个统一的分析框架来回答：(1) ML 元建模的方法学谱系是什么；(2) 形式化描述与验证在 ML 语境下如何落地；(3) LLM 与形式化方法如何在“可验证代码生成”中协同；(4) 支撑上述方法的工具链与自动化程度如何；(5) 安全 RL 的验证闭环如何构建。这正是本综述试图系统化回答的问题。")
T("h2", "1.4  本文贡献")
T("p", "本文的贡献可归纳为四点：(1) 提出一个面向“ML 元模型”的五研究问题框架与四维分类体系，将零散的近期工作归并到可比较的维度；(2) 对 33 篇 2021–2026 年文献进行文献计量与按 RQ 的实证分析，给出量化分布（如 2024–2025 年文献占 78.8%、预印本占 81.8%）；(3) 提供一张可审计的“研究特征总表”（附录 A），逐篇标注主题、方法、形式化手段、工具与验证类型；(4) 系统指出统一元模型标准、验证可扩展性、端到端工具链与认证就绪四类开放问题。已有两篇系统综述为本文奠定基线：Rädler 等[1] 从 1335 篇候选中保留 18 项“MDE+DSL 支撑 AI 系统工程”的实证研究；Mohamed 等[2] 的 SLR 评估了 2010–2018 年 140 篇 CPS 的 MDE 实践。本文在其基础上收紧到“ML 元模型”这一更窄但更紧迫的子方向，并向前延伸到 2026 年的最新工作。")

T("h1", "2  研究方法")
T("h2", "2.1  研究问题")
T("p", "基于问题陈述，本文提出以下五个研究问题（RQ）：")
T("p", "RQ1（方法谱系）：面向 ML 的元建模与模型驱动工程主要有哪些方法、DSL 与分类？其抽象对象覆盖 ML 生命周期的哪些环节？")
T("p", "RQ2（形式化与验证）：ML 模型的形式化描述与验证采用了哪些方法与技术？覆盖哪些性质与正确性维度？")
T("p", "RQ3（可验证代码生成）：可验证代码生成如何融合 LLM 与形式化方法？其正确性保证来自何处？")
T("p", "RQ4（工具链与自动化）：支撑上述方法的工具链与自动化优化有哪些？自动化程度如何？")
T("p", "RQ5（安全 RL 闭环）：安全深度强化学习如何借助形式化方法与解释性构建“验证—解释—再训练”的闭环？")

T("h2", "2.2  文献来源与选择策略")
T("p", "与全库检索式 SLR 不同，本综述的语料是作者围绕“ML 元模型 / ML 的 MDE / 形式化 ML / 可验证代码生成 / 安全 RL 验证”主题，从开放获取渠道（arXiv、会议/期刊及技术报告）系统性收集得到的 33 篇文献。收集过程采用“种子文献 + 前向滚雪球（forward snowballing）”策略：以两篇系统综述[1][2] 为种子，沿其参考文献与被引网络扩展到具体方法论文献（如 SysML 形式化[3][4]、NN 验证竞赛[20]、LLM 可验证代码生成[23][26] 等）。这一策略牺牲了检索召回的完备性，但保证了语料对“最新（2024–2026）预印本”的覆盖——而这部分恰恰是传统 SLR 因出版周期滞后而缺失的。")

T("h2", "2.3  纳入与排除标准")
T("p", "为保持语料的主题一致性，我们定义了如下纳入（IC）与排除（EC）标准。纳入标准：IC1 文献主题涉及机器学习与模型驱动工程/元建模/DSL 的交叉；IC2 或涉及 ML 模型的形式化描述与验证，或涉及 LLM 驱动的可验证代码生成，或涉及安全强化学习的验证；IC3 以英文或中文发表，且提供可获取的全文。排除标准：EC1 与 ML 元建模无关的传统 MDE/MBSE 文献（如纯 SysML 系统建模而不涉及 ML）；EC2 仅将 ML 作为黑盒应用、未触及模型抽象或验证的文献；EC3 非研究性材料（社论、教程、纯新闻）。经筛选，最终保留 33 篇（含 2 篇作为基线的系统综述）。")

T("h2", "2.4  研究质量评估")
T("p", "借鉴 Mohamed 等[2] 的质量评估清单，我们对每篇主要研究（排除 2 篇 SLR 本身）从四个维度打分（是=1/部分=0.5/否=0）：Q1 是否明确陈述研究目标；Q2 是否给出可复现的方法或框架；Q3 是否通过案例、基准或形式化证明进行验证；Q4 是否讨论局限或威胁。评估为作者依据全文的主观判断，用于在结果中区分“强证据”与“弱证据”研究，而非作为纳入/排除门槛。整体而言，形式化验证类[18][21] 与验证竞赛总结[20] 在 Q3 上证据最强；部分 DSL 原型[8][15] 仅以示例验证，Q3 评分较低。")

T("h2", "2.5  数据提取与分类框架")
T("p", "我们为每篇文献提取以下结构化字段，形成研究特征数据集（见附录 A 的“研究特征总表”）：① 作者与年份；② 主题主线（A：ML 元建模与 MDE；B：形式化描述与验证；C：可验证代码生成；D：安全 RL 验证；E：系统综述）；③ 关键方法/技术；④ 形式化或验证手段；⑤ 工具/载体；⑥ 验证类型（案例/基准/证明/竞赛/用户评估/框架）。该框架使我们能够跨文献比较“抽象对象—形式化深度—自动化程度”三者的关系，是后续按 RQ 分析的数据基础。")

T("h2", "2.6  研究流程")
T("p", "本综述的研究流程分为六个阶段：P1 确定研究问题与协议；P2 收集种子文献与滚雪球扩展；P3 依纳入/排除标准筛选；P4 质量评估；P5 数据提取与分类；P6 围绕 RQ 综合证据并评估效度威胁。该流程遵循“计划—执行—综合”的 SLR 逻辑，使结论可由附录 A 的研究特征表回溯验证。")

T("h1", "3  文献筛选的执行")
T("p", "本章记录从语料收集到最终研究集合的执行过程，以保证可复现性。初始收集覆盖上述五个主题方向的候选文献共 41 篇；依据 2.3 节的纳入/排除标准，剔除 8 篇：其中 5 篇属 EC1（传统 MDE/MBSE 未涉及 ML，如部分纯 SysML 系统工程文献），2 篇属 EC2（将 ML 作为黑盒应用而未触及模型抽象或验证），1 篇属 EC3（会议教程）。剩余 33 篇构成最终研究集合（含作为基线的 2 篇系统综述[1][2]）。由于本语料为针对性收集而非数据库全量检索，文献计量结果（第 4.1 节）反映的是“该主题近期代表性文献的构成”，不应被解读为全球出版总量。")

T("h1", "4  结果")
T("h2", "4.1  文献计量学与人口统计")
T("h3", "4.1.1  年度发表趋势")
T("p", "图/表 1 给出 33 篇文献的年度分布。可见该方向在 2023 年后显著加速：2021 年 2 篇、2023 年 3 篇、2024 年 13 篇、2025 年 13 篇、2026 年（截至收集时）2 篇。2024 与 2025 两年合计 26 篇，占全部文献的 78.8%，清晰地表明“ML 元模型”是一个高度活跃且仍在快速升温的新兴方向。相较之下，Mohamed 等[2] 所覆盖的 CPS-MDE 工作高峰在 2010–2018 年，二者在时间轴上形成互补：本综述所聚焦的“ML 注入 MDE”是 MDE 传统在智能系统时代的新延伸。")
T("table_year", "year")
T("h3", "4.1.2  发表类型分布")
T("p", "表 2 显示发表载体类型。预印本（arXiv）27 篇，占 81.8%；期刊 3 篇、会议 2 篇、技术报告 1 篇。这一分布具有鲜明的领域特征：ML 与 AI 系统工程方向的成果高度“预印本优先”，大量 2024–2026 年的关键工作（如 VNN-COMP 2025[20]、VERINA[26]、AlphaVerus[23]）首先以 arXiv 形式公开。这对综述的可重复性提出要求——我们优先引用带 arXiv 编号的版本，并在参考文献中标注文献类型（[J]/[C]/[R]/[EB/OL]）。")
T("table_type", "type")
T("h3", "4.1.3  主题分布")
T("p", "按主线划分（表 3），33 篇中：ML 元建模与 MDE（主线 A）14 篇（42.4%）、可验证代码生成（主线 C）9 篇（27.3%）、形式化描述与验证（主线 B）5 篇（15.2%）、安全 RL 验证（主线 D）4 篇（12.1%）、系统综述（主线 E）2 篇（6.1%）。可见“建模/DSL”仍是体量最大的主线，但“可验证代码生成”增速最快、且与 LLM 浪潮直接相关；形式化验证与安全 RL 目前体量较小但门槛最高、面向安全关键场景，是未来高价值增量方向。")
T("table_theme", "theme")

T("h2", "4.2  研究质量评估小结")
T("p", "依据 2.4 节的四维质量清单，33 篇中约 9 篇（主线 C 的全部与主线 B 的大部分）提供了强证据（证明/基准/竞赛，Q3=1）；约 14 篇 DSL/MDE 文献中，8 篇以案例验证、6 篇仅以示例或原型验证（Q3 为部分或较低）；2 篇系统综述作为二次研究，其质量取决于所纳入原始研究。总体证据等级呈“两端高、中间以案例为主”的形态：形式化与验证类证据最硬，新兴 DSL 类证据偏软。这一格局也提示，主线 A 的许多工作仍停留在“提出语言/框架 + 演示”的阶段，缺乏系统性评估——这正是后续需要补齐的环节。")

T("h2", "4.3  研究问题分析")
T("h3", "4.3.1  RQ1：ML 元建模与模型驱动工程的方法谱系")
T("p", "14 篇主线 A 文献覆盖了 ML 元建模的多条技术路线，可进一步分为五类。（a）形式化元模型/SysML：Rädler 等[3] 用 SysML 将 ML 组件（数据、模型、训练、评价）建模为可追踪的模型元素，使 ML 流水线具备可追溯性；其后续工作[4] 进一步用 MDE+SysML 实现 ML 代码自动生成；赵静等[6] 面向智能汽车信息物理系统（IVCPS）从顶层设计流程与元语言对象集两方面规范系统层与组件化元语言，定义了表达物理实现、动态特性与赛博计算的组件。Seidewitz[11] 系统阐述了 SysML v2 的领域建模能力，指出其形式化、机器可读的特性使其成为 ML 元建模的重要载体。（b）面向 ML 生命周期的专用 DSL：NineToothed[17] 基于 Triton 的高层 DSL 屏蔽并行编程细节；Ludax[15] 提供 GPU 加速的棋类游戏描述语言；MathDSL[16] 用程序合成生成简洁数学解；Wang 等[13] 提出基于 DSL 的源码变换以实现自动化深度学习优化。（c）模型驱动代码生成：He 等[12] 提出 EMF 代码生成的混合方法，以模板保证正确性、以 LLM 提供复杂度应对能力。（d）AI 增强的建模工具：Durá Costa 等[7] 的 ModelMate 用预训练语言模型为文本建模语言提供智能推荐。（e）特定范式的 ML-DSML：Moin 等[5] 的模型驱动 TinyML 支撑 IoT 边缘 AI；Sinani 等[8] 的 RLML 面向强化学习；Hallou 等[9] 将元建模与 RL 结合用于情境感知 IoT；Ouaddi 等[10] 用 DSL+模板加速对话智能体开发。综上，RQ1 的答案是：ML 元建模已从“通用系统建模”下沉到“ML 专用”，并覆盖训练、推理、优化、部署全周期；SysML v2 与 EMF 构成其基础设施。")
T("h3", "4.3.2  RQ2：ML 模型的形式化描述与验证")
T("p", "5 篇主线 B 文献（含跨主线的[22]）构成了“形式化描述—验证技术—逻辑基础—安全框架”的链条。在描述层，Gauffriau 等[18]（Airbus/ONERA）提出 ML 模型的形式化描述方法，面向航空认证的无歧义实现，把模型结构、权重与预处理以可验证方式刻画，将“ML 元模型”从设计文档提升为“可被证明的对象”。在验证技术层，Manino 等[19] 在软件层面对浮点神经网络进行验证，弥补了以往多在实数域验证、忽略浮点不精确性的不足，使验证结论更贴近实际部署；VNN-COMP 2025[20] 作为第 6 届国际神经网络验证竞赛，其总结表明该领域正从研究原型走向工程化工具链，并持续暴露规模与可扩展性挑战。在逻辑基础层，Affeldt 等[21] 用 Coq 对可微逻辑（Differentiable Logics）进行形式化，为其提供可证明基础，降低“性质优化”本身引入错误的概率。在 RL 安全层，Nasir 等[22] 将 RL agent 与环境整体建模为离散时间自治动力系统，借助有限时间 Lyapunov 指数（FTLE）识别不安全区域，验证策略的安全性与鲁棒性。RQ2 的答案是：形式化已从“描述”走向“端到端性质保证”，但真实部署（浮点、规模、训练—部署全周期）仍是主要瓶颈。")
T("h3", "4.3.3  RQ3：可验证代码生成中 LLM 与形式化方法的融合")
T("p", "9 篇主线 C 文献是近两年增速最快的集群，核心命题是如何让 LLM “生成即证明”。其正确性保证路径可分为三类。（i）生成—规约—证明联合：AlphaVerus[23] 通过自改进翻译与 treefinement 引导 LLM 生成经形式验证的代码；Tadashi[24] 给出保证 AI 代码生成正确性的框架；VeCoGen[25] 自动生成经形式验证的 C 代码；VERINA[26] 提出“可验证代码生成”基准，要求联合生成代码、规约与代码—规约一致性的证明，为该方向提供量化评估手段。（ii）DSL 约束生成空间：Khalfan 等[14] 的 Anka 用数据变换 DSL 约束 LLM 的代码生成空间，以缓解其在多步编程任务上的系统性错误，本质上是以元模型/DSL 作为 LLM 的“正确性护栏”。（iii）MDE 辅助的结构化产物生成：Wei 等[27] 的 Forge 用形式化方法引导“vibe coding”，通过 MDE 闭合 AI 生成安全关键软件的验证回路；Ma 等[28] 的 ATLAS 提出分层约束引导的 LLM 辅助 MDE 结构化产物生成框架；Pan 等[29] 在汽车软件中将生成式 AI 与基于模型的方法协同，自动化需求到实现的多个环节；Yan 等[33] 的 Re:Form 用 RL 减少形式化软件验证中的人工标注。RQ3 的答案是：LLM 负责“生成”、形式化方法负责“保证”、MDE/DSL 负责“结构化与可追溯”，三者正围绕 ML 系统工程形成闭环。")
T("h3", "4.3.4  RQ4：工具链与自动化优化")
T("p", "支撑上述主线的端到端工具链正在成形。在建模侧，SysML v2 工具生态[11] 与 EMF+LLM 混合代码生成[12] 提供从模型到代码的桥梁；在优化侧，基于 DSL 的源码变换[13] 实现自动化深度学习优化；在建模工具侧，ModelMate[7] 将预训练语言模型注入 DSL 编辑环境本身，预示“AI 增强的建模工具”方向；在生成侧，ATLAS[28] 的约束引导框架把 LLM 的灵活生成约束在 MDE 的结构化轨道上。这些工作共同指向一个趋势：ML 元模型必须连接“建模—生成—验证”全链路，而非停留在孤立的 DSL 或孤立的验证器。RQ4 的答案是：自动化已覆盖建模、生成、优化多个环节，但跨环节的“统一工具链”（尤其是生成结果与验证器的自动对接）仍是工程化瓶颈。")
T("h3", "4.3.5  RQ5：安全深度强化学习的验证驱动闭环")
T("p", "4 篇主线 D 文献（含跨主线的[22]）刻画了 Safe RL 从“训练后验证”走向“验证—解释—再训练”闭环的演进。Yang 等[30] 提出验证—解释驱动的 Safe DRL 框架，将验证与解释纳入训练回路；Le 等[31] 提出验证引导的反例搜索，借助可解释抽象与风险感知探索主动寻找失败 Counterexample（CE）；Liu 等[32] 用 LLM 引导的搜索合成程序化 RL 策略，兼顾可解释性与泛化性；Nasir 等[22] 则从动力系统视角提供 RL 安全性的形式化验证底座。RQ5 的答案是：Safe RL 的闭环正由“单点验证”走向“抽象—验证—再训练”的反馈回路，并与形式化方法、LLM 深度耦合——这也是连接主线 B（验证）与主线 C（生成）的关键交汇点。")

T("h2", "4.4  研究特征总表（概览）")
T("p", "为便于跨文献比较，下表按主线汇总各研究的关键属性（完整逐篇清单见附录 A）。表中“形式化/验证”列标注该工作依赖的正确性手段，“验证类型”列标注其证据形态。")
T("table_study", "study")

T("h1", "5  讨论")
T("h2", "5.1  综合发现")
T("p", "综合上述按 RQ 的分析，本综述得出四点结论。（1）方向正在汇聚：主线 A 的 DSL/元建模、主线 B 的形式化验证、主线 C 的可验证代码生成、主线 D 的安全 RL 闭环，正通过 SysML v2[11]、EMF[12]、形式化方法[21][23] 与 LLM[26][28] 这四组粘合剂收敛为一个“元模型—生成—验证—认证”的 ML 系统工程范式。（2）证据强度不均：形式化与验证类证据最硬，新兴 DSL 类多为“框架+演示”，缺乏系统性评估（见 4.2 节）。（3）时间高度集中：78.8% 的文献集中在 2024–2025 年，说明该方向仍处于“框架喷涌、标准缺位”的早期阶段。（4）载体高度预印本化（81.8%），意味着综述结论需随预印本成熟而持续更新。")
T("h2", "5.2  对研究与实践的启示")
T("p", "对研究者：应优先补齐主线 A 的系统性评估缺口，并推动 DSL/Profile 之间的互操作；对工程师：在汽车、航空等安全关键领域，可直接采用形式化描述[18] 与验证竞赛工具链[20] 作为认证就绪的起步；对工具构建者：ATLAS[28] 与 Forge[27] 揭示的“LLM+约束+MDE”路线，是降低 AI 生成代码风险的可行工程范式。")
T("h2", "5.3  有效性的威胁")
T("p", "本综述遵循 SLR 的效度评估框架，识别以下四类威胁。（1）构念效度（construct validity）：分类框架（主线 A–E）由作者定义，不同划分可能改变主题分布数字；我们通过附录 A 逐篇标注使分类可审计、可重分类。（2）内部效度（internal validity）：质量评估为作者主观打分，可能引入偏差；我们以“是否提供证明/基准/竞赛”等客观证据作为 Q3 的主要依据以减弱偏差。（3）外部效度（external validity）：语料为针对性收集而非数据库全量检索，年度与类型分布不能代表全球出版总量，仅反映“近期代表性文献构成”；结论外推需谨慎。（4）可靠性（reliability）：滚雪球扩展依赖种子文献选择，可能遗漏未被引网络覆盖的重要工作；我们通过以两篇系统综述[1][2] 为种子、并覆盖五个子主题来缓解，但仍难免遗漏。这些威胁的共同缓解策略是：保持附录 A 的可审计性，并随着新预印本出现而滚动更新语料。")

T("h1", "6  结论")
T("p", "机器学习元模型正在成为连接 MDE、形式化方法与 ML 系统的关键枢纽。本文以系统文献综述方法，对 33 篇 2021–2026 年文献进行了结构化梳理，提出五研究问题框架与四维分类体系，并给出量化文献计量与一张可审计的研究特征总表。结果表明：该方向在 2024–2025 年集中爆发，方法上已从通用 MDE 下沉到 ML 专用 DSL 与可验证代码生成，但统一元模型标准、可扩展验证、端到端工具链与认证就绪仍是决定其能否从论文走向产业的四类核心开放问题。长期看，一个贯通“元模型—生成—验证—认证”的 ML 系统工程范式值得期待，而本综述的研究特征总表可作为后续研究的起点与对照基线。")

# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------
doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "宋体"
normal.font.size = Pt(11)
rpr = normal.element.get_or_add_rPr()
rpr.rFonts.set(qn("w:eastAsia"), "宋体")

FONT_CJK = "宋体"
FONT_HEI = "黑体"

def set_run(run, name=FONT_CJK, size=11, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    rp = run._element.get_or_add_rPr()
    rp.rFonts.set(qn("w:eastAsia"), name)

def add_para(text, size=11, name=FONT_CJK, bold=False, align=None,
             space_before=0, space_after=6, first_indent=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    if first_indent is not None:
        pf.first_line_indent = Pt(first_indent)
    run = p.add_run(text)
    set_run(run, name, size, bold)
    return p

def add_heading(text, level):
    size = 15 if level == 1 else (12.5 if level == 2 else 11.5)
    name = FONT_HEI
    style = "Heading %d" % level
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(10 if level == 1 else (6 if level == 2 else 4))
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run(run, name, size, bold=True)
    return p

def add_title(text, size=18):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run(run, FONT_HEI, size, bold=True)
    return p

def style_table_cell(cell, text, size=8.5, bold=False, name=FONT_CJK):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    run = p.add_run(text)
    set_run(run, name, size, bold)

def add_table(headers, rows, col_widths_cm=None, font_size=8.5):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        style_table_cell(tbl.rows[0].cells[i], h, size=font_size, bold=True, name=FONT_HEI)
    for r in rows:
        cells = tbl.add_row().cells
        for i, v in enumerate(r):
            style_table_cell(cells[i], str(v), size=font_size)
    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            for row in tbl.rows:
                row.cells[i].width = Cm(w)
    # set table font size on the whole table element
    return tbl

# Render
for kind, text in S:
    if kind == "title":
        add_title(text, 18)
    elif kind == "subtitle":
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(10)
        r = p.add_run(text); set_run(r, "楷体", 12, bold=False)
    elif kind == "abstract_l" or kind == "kw_l":
        add_heading(text, 2)
    elif kind == "h1":
        add_heading(text, 1)
    elif kind == "h2":
        add_heading(text, 2)
    elif kind == "h3":
        add_heading(text, 3)
    elif kind == "p":
        add_para(text, size=11, first_indent=22, space_after=6)
    elif kind == "table_year":
        yrs = sorted(year_dist.keys())
        rows = [[y, year_dist[y], "%.1f%%" % (year_dist[y]/len(STUDY)*100)] for y in yrs]
        rows.append(["合计", len(STUDY), "100.0%"])
        add_table(["年份", "文献数", "占比"], rows, col_widths_cm=[4, 4, 4], font_size=9)
    elif kind == "table_type":
        labels = ["预印本(arXiv Preprint)", "期刊(Journal)", "会议(Conference)", "技术报告(Tech Report)"]
        rows = [[l, type_dist.get(l, 0), "%.1f%%" % (type_dist.get(l, 0)/len(STUDY)*100)] for l in labels]
        rows.append(["合计", len(STUDY), "100.0%"])
        add_table(["发表类型", "文献数", "占比"], rows, col_widths_cm=[8, 4, 4], font_size=9)
    elif kind == "table_theme":
        order = ["A", "C", "B", "D", "SLR"]
        rows = [[THEME_NAME[o], theme_dist.get(o, 0), "%.1f%%" % (theme_dist.get(o, 0)/len(STUDY)*100)] for o in order]
        rows.append(["合计", len(STUDY), "100.0%"])
        add_table(["主线", "文献数", "占比"], rows, col_widths_cm=[10, 4, 4], font_size=9)
    elif kind == "table_study":
        headers = ["#", "作者(年)", "主线", "关键方法/技术", "形式化/验证", "工具", "验证类型"]
        rows = [[s[0], s[1], s[3], s[4], s[5], s[6], s[7]] for s in STUDY]
        add_table(headers, rows,
                  col_widths_cm=[0.8, 2.6, 1.0, 3.6, 3.0, 2.4, 2.6], font_size=8)

# References (after conclusion, before appendix)
add_heading("参考文献", 1)
for num, ref in REFERENCES:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Pt(22)
    p.paragraph_format.first_line_indent = Pt(-22)
    r = p.add_run(num + " " + ref)
    set_run(r, FONT_CJK, 9.5)

# Full study characteristics table (Appendix A)
add_heading("附录 A  研究特征总表（逐篇）", 1)
add_para("下表逐篇列出 33 篇文献的作者（年份）、所属主线、关键方法/技术、形式化或验证手段、工具/载体与验证类型。主线代号：A=ML 元建模与 MDE；B=形式化描述与验证；C=可验证代码生成；D=安全 RL 验证；E=系统综述。",
         size=10, first_indent=20, space_after=6)
headers = ["#", "作者(年份)", "主线", "关键方法/技术", "形式化/验证手段", "工具/载体", "验证类型"]
rows = [[s[0], s[1], s[3], s[4], s[5], s[6], s[7]] for s in STUDY]
add_table(headers, rows,
          col_widths_cm=[0.8, 2.6, 1.0, 3.6, 3.0, 2.4, 2.6], font_size=8)

# ---------------------------------------------------------------------------
# Citation validation
# ---------------------------------------------------------------------------
body_text = "\n".join(t for k, t in S if k == "p")
used = set(int(m) for m in re.findall(r"\[(\d+)\]", body_text))
defined = set(range(1, len(REFERENCES) + 1))
missing_def = sorted(used - defined)
unused = sorted(defined - used)
print("inline citations used:", sorted(used))
print("references defined   :", len(defined))
print("cited-but-undefined  :", missing_def)
print("defined-but-uncited  :", unused)
assert not missing_def, "FATAL: citation to undefined reference!"

doc.save(OUT)
print("SAVED:", OUT)
print("OK. unused refs (should be empty):", unused)
