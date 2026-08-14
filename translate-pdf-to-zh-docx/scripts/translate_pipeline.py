"""Translate academic PDFs (English or other) into clean, structured Chinese Word docs.

Cleanup pipeline (v4 — general, applies to every paper):
  * Header/footer / running-head removal (position bands + cross-page repeats
    + explicit meta phrases + footnote symbols).
  * Paragraph re-assembly BEFORE translation: cross-block / cross-page sentence
    splits are healed so whole paragraphs translate in context. Two-column
    reading order is reconstructed so a left-column paragraph is never broken
    by the right column.
  * Figures: every embedded raster image is cropped from its real page position
    and embedded in reading order (tiny decorations skipped).
  * Tables: cells are extracted with PyMuPDF find_tables().extract(), each cell
    is translated, and a real Word table is produced (never a garbled image).
  * Math: spans are detected by math font or math-glyph density. A line that is
    essentially all-math (no readable prose) is rendered as a clean centered
    image at its natural width (capped to the page, never oversized); mixed
    "text + formula" lines stay as text so surrounding words are never captured.
  * Footnote / reference URLs collected once under a "参考链接" appendix.
  * CJK typography: body 宋体 + Times New Roman, headings 黑体 + Arial,
    first-line indent, 1.5 line spacing.

Usage:
  translate_pipeline.py [--force] <pdf_or_dir> [pdf_or_dir ...]
  Output: <basename>-翻译.docx next to each source PDF.
  --force : re-translate even if the output already exists.
"""
import io
import fitz, re, json, time, os, sys, argparse, urllib.parse, urllib.request, ssl, collections
from concurrent.futures import ThreadPoolExecutor
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

_CTRL = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")
_URL = re.compile(r"https?://[^\s\u4e00-\u9fff]+", re.I)
_PURE_NUM = re.compile(r"^\d{1,4}$")
_SECTION = re.compile(r"^\d+(\.\d+)*\.?\s+[\u4e00-\u9fffA-Za-z]")
_SECTION_SPLIT = re.compile(
    r"^(\d+(?:\.\d+)*\.?\s+(?:[A-Za-z][\w-]*\s+){1,2})(?=[A-Za-z])")
_RUNHEAD = re.compile(
    r"manuscript submitted|published as a conference paper|preprint|"
    r"©|copyright|all rights reserved|under review|arxiv:\d|doi:|"
    r"available at|submitted to|received|accepted|academic editor|"
    r"correspondence|to appear in|in press|under consideration|preprint submitted|"
    r"this is a preprint|version of record|contains two columns|journal of|"
    r"springer|ieee transactions|acm transactions|proceedings of|conference on",
    re.I)
_FOOTNOTE_MARK = re.compile(r"^[⋆★☆†‡¹²³⨯✱*]+")
_FOOTNOTE_TAIL = re.compile(r"[⋆★☆†‡¹²³⨯✱*]+$")
_LIST_ITEM = re.compile(r"^(\d+[\.\)]|[\u2022\-*]\s|[a-z]\)|[ivx]+\.)\s", re.I)
_END_PUNCT = set(".!?;:。！？；：)]}\"')»”’")
_MATH_FONT = re.compile(
    r"cmmi|cmsy|cmex|stix|cambria\s*math|latin\s*modern\s*math|mathjax|"
    r"ams\s*math|rsfs|euler|symbol|math\b", re.I)
_MATH_UNI = set(
    "∈∉⊂⊃⊆⊇∪∩∀∃→↔⇒⇐↦≤≥×÷√∑∏∫∂∇∞≈≠≡±∓∝"
    "αβγδεζηθικλμνξπρςστυφχψωΑΒΓΔΘΛΞΠΣΦΨΩ"
    "ℓℜℑ⊥∥∠∧∨¬◊•†‡ℵℏ℘⊙⊗⊕∆"
    "′″‴⟨⟩⟪⟫⌊⌋⌈⌉〈〉")

BODY_CJK = "宋体"
BODY_LATIN = "Times New Roman"
HEAD_CJK = "黑体"
HEAD_LATIN = "Arial"
BODY_SIZE = 11
SENT = "␞M%d␞"
_SENT_RE = re.compile(r"␞M(\d+)␞")

PROXY = os.environ.get("TRANSLATE_PROXY", "http://127.0.0.1:7897")
OUT_EXT = ".docx"
BASE_GT = "https://translate.google.com/translate_a/single"
UA = ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
      "(KHTML, like Gecko) Chrome/124.0 Safari/537.36")
ctx = ssl.create_default_context()
ctx.check_hostname = False
ctx.verify_mode = ssl.CERT_NONE
_op = urllib.request.build_opener(
    urllib.request.ProxyHandler({"https": PROXY, "http": PROXY}),
    urllib.request.HTTPSHandler(context=ctx),
)
# Direct opener with NO proxy, used for the MyMemory fallback so it is not
# routed through the (possibly dead) local proxy.
_DIRECT = urllib.request.build_opener(urllib.request.ProxyHandler({}))


def clean(s):
    return _CTRL.sub("", s) if s else s


def set_run_font(run, latin, cjk, size=None):
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), cjk)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    if size:
        run.font.size = Pt(size)


def set_cjk(style, latin, cjk, size=None):
    style.font.name = latin
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), cjk)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    if size:
        style.font.size = Pt(size)


_PROXY_DEAD = False


def _translate_one(text, sl="auto", tl="zh-CN"):
    """Try Google (via the local proxy). If the proxy is unreachable, fall back
    to MyMemory (direct, no proxy) which is reachabe from China. Either failure
    returns the original text so the run never hard-crashes."""
    global _PROXY_DEAD
    if not _PROXY_DEAD:
        try:
            q = urllib.parse.quote(text)
            url = "%s?client=gtx&sl=%s&tl=%s&dt=t&q=%s" % (BASE_GT, sl, tl, q)
            req = urllib.request.Request(url, headers={"User-Agent": UA})
            data = _op.open(req, timeout=30).read().decode("utf-8", "ignore")
            j = json.loads(data)
            return "".join(s[0] for s in j[0] if s and s[0])
        except Exception as e:
            _PROXY_DEAD = True
            sys.stderr.write("  google proxy unreachable -> MyMemory fallback: %s\n"
                             % str(e)[:60])
    try:
        src = sl if sl not in ("auto", "") else "en"
        # MyMemory public API limits each query to ~500 chars; split long text
        # and translate piece by piece. Any error/limit message is treated as a
        # failure so the original text is kept (never leaks garbage).
        pieces = re.split(r'(?<=[.!?;:，。！？；：])\s*', text)
        buf, segs = "", []
        for p in pieces:
            if len(buf) + len(p) <= 450:
                buf += p + " "
            else:
                if buf:
                    segs.append(buf.strip())
                buf = p + " "
        if buf:
            segs.append(buf.strip())
        if not segs:
            segs = [text]
        out_parts = []
        for seg in segs:
            q = urllib.parse.quote(seg)
            url = ("https://api.mymemory.translated.net/get?q=%s&langpair=%s|%s"
                   % (q, src, tl))
            req = urllib.request.Request(url, headers={"User-Agent": UA})
            data = _DIRECT.open(req, timeout=20).read().decode("utf-8", "ignore")
            j = json.loads(data)
            t = j.get("responseData", {}).get("translatedText", "")
            if (not t or "MYMEMORY WARNING" in t or "QUOTA" in t.upper()
                    or "QUERY LENGTH LIMIT" in t or "MAX ALLOWED" in t
                    or "EXCEEDED" in t.upper()):
                out_parts.append(seg)
            else:
                out_parts.append(t)
        return " ".join(out_parts)
    except Exception:
        return text


def gtranslate(text, sl="auto", tl="zh-CN", retries=4):
    text = text.strip()
    if not text:
        return text
    if len(text) <= 4500:
        chunks = [text]
    else:
        parts = re.split(r'(?<=[.!?;:])\s+', text)
        chunks, buf = [], ""
        for p in parts:
            if len(buf) + len(p) < 4500:
                buf += p + " "
            else:
                if buf:
                    chunks.append(buf.strip())
                buf = p + " "
        if buf:
            chunks.append(buf.strip())
    out = []
    for ch in chunks:
        out.append(_translate_one(ch, sl, tl))
    return "\n".join(out)


def in_rect(x0, y0, x1, y1, rects):
    cx, cy = (x0 + x1) / 2, (y0 + y1) / 2
    for r in rects:
        if r[0] <= cx <= r[2] and r[1] <= cy <= r[3]:
            return True
    return False


# ---------------------------------------------------------------------------
# Math detection
# ---------------------------------------------------------------------------
def span_math(sp):
    if _MATH_FONT.search(sp.get("font", "")):
        return True
    t = sp.get("text", "")
    if not t:
        return False
    mc = sum(1 for c in t if c in _MATH_UNI)
    return len(t) > 0 and mc / len(t) >= 0.3


def line_is_display(line):
    """A line is a display/standalone equation ONLY when it is essentially all
    math (no readable prose). Mixed 'text + formula' lines are kept as text so
    the surrounding words are never captured into the equation image.

    Detection:
      * split spans into math vs text;
      * if any TEXT span contains a readable word (>=2 CJK chars, or a run of
        >=3 ASCII letters) the line is prose -> not display;
      * otherwise a near-pure-math line (>=2 math spans AND >=50% math spans,
        or very high math-glyph density) is a display equation.
    """
    spans = [s for s in line if s.get("text", "").strip()]
    if not spans:
        return False
    math_spans = [s for s in spans if span_math(s)]
    for s in spans:
        if span_math(s):
            continue
        t = s.get("text", "")
        if re.search(r"[\u4e00-\u9fff]{2,}", t):
            return False
        if re.search(r"[A-Za-z]{3,}", t):
            return False
    if len(math_spans) >= 2 and len(math_spans) / len(spans) >= 0.5:
        return True
    txt = "".join(s["text"] for s in spans)
    if txt and sum(1 for c in txt if c in _MATH_UNI) / len(txt) >= 0.6:
        return True
    return False


def union_bbox(spans):
    return (min(s["bbox"][0] for s in spans), min(s["bbox"][1] for s in spans),
            max(s["bbox"][2] for s in spans), max(s["bbox"][3] for s in spans))


_FOOTNOTE_ANY = re.compile(r"[⋆★☆†‡¹²³⨯✱*]")


def prep_line(line):
    # Inline math is kept as translated text (cheap, usually readable).
    # Only display equations (math-dominant lines) are rendered as images.
    # Footnote markers (stars/daggers/superscript digits) never appear in real
    # content, so strip them from every span up front.
    parts = []
    for sp in line:
        t = _FOOTNOTE_ANY.sub("", sp.get("text", "")).strip()
        if t:
            parts.append(t)
    return " ".join(parts), []


# ---------------------------------------------------------------------------
# Cropping / rendering
# ---------------------------------------------------------------------------
_PDF_CACHE = {}


def _get_doc(pdf_path):
    if pdf_path not in _PDF_CACHE:
        _PDF_CACHE[pdf_path] = fitz.open(pdf_path)
    return _PDF_CACHE[pdf_path]


def render_crop(pdf_path, page_index, bbox, zoom=2.2):
    pdoc = _get_doc(pdf_path)
    pg = pdoc[page_index]
    clip = fitz.Rect(bbox)
    pix = pg.get_pixmap(clip=clip, matrix=fitz.Matrix(zoom, zoom))
    return io.BytesIO(pix.tobytes("png")), pix.width, pix.height


# ---------------------------------------------------------------------------
# Page element extraction
# ---------------------------------------------------------------------------
def order_elements(elems, W):
    xs = [(e["bbox"][0] + e["bbox"][2]) / 2 for e in elems]
    split = None
    if len(xs) >= 4:
        lo, hi = W * 0.18, W * 0.82
        nb = 20
        bins = [0] * nb
        for xc in xs:
            if lo <= xc <= hi:
                b = int((xc - lo) / (hi - lo) * (nb - 1))
                bins[b] += 1
        best = cur = beststart = start = 0
        for i, v in enumerate(bins):
            if v == 0:
                cur += 1
            else:
                if cur > best:
                    best = cur
                    beststart = start
                cur = 0
                start = i + 1
        if cur > best:
            best = cur
            beststart = start
        if best >= 4:
            split = lo + (beststart + best / 2) / nb * (hi - lo)
    for e in elems:
        xc = (e["bbox"][0] + e["bbox"][2]) / 2
        e["_col"] = 0 if split is None or xc < split else 1
    elems.sort(key=lambda e: (e["_col"], round(e["y0"] / 6), e["bbox"][0]))
    return elems


def extract_page(page, W, H):
    d = page.get_text("dict")
    table_rects = []
    table_data = []
    try:
        tf = page.find_tables()
        for t in tf.tables:
            try:
                nrows = len(t.rows)
                ncols = len(t.cols)
            except Exception:
                try:
                    ex = t.extract()
                    nrows = len(ex)
                    ncols = len(ex[0]) if ex else 0
                except Exception:
                    nrows = ncols = 0
            if 2 <= nrows <= 60 and 2 <= ncols <= 12 and nrows * ncols >= 4:
                bbox = tuple(t.bbox)
                table_rects.append(bbox)
                cells = None
                try:
                    cells = t.extract()
                except Exception:
                    cells = None
                table_data.append((bbox, cells))
    except Exception:
        pass

    img_els = []
    try:
        iinfo = page.get_image_info(xrefs=True)
        for it in iinfo:
            r = it.get("bbox")
            if not r:
                continue
            x0, y0, x1, y1 = r
            if min(x1 - x0, y1 - y0) < 12:
                continue
            if in_rect(x0, y0, x1, y1, table_rects):
                continue
            img_els.append({"kind": "image", "bbox": (x0, y0, x1, y1),
                            "page": page.number, "y0": y0, "y1": y1,
                            "x0": x0})
    except Exception:
        pass

    text_els = []
    for b in d["blocks"]:
        if b.get("type", 0) != 0:
            continue
        x0, y0, x1, y1 = b["bbox"]
        if in_rect(x0, y0, x1, y1, table_rects):
            continue
        lines = []
        sizes = []
        for line in b["lines"]:
            spans = []
            for s in line["spans"]:
                txt = s["text"]
                if not txt.strip():
                    continue
                spans.append({"text": txt, "font": s["font"], "size": s["size"],
                              "bbox": tuple(s["bbox"])})
                sizes.append(s["size"])
            if spans:
                lines.append(spans)
        if not lines:
            continue
        text = " ".join(" ".join(sp["text"] for sp in ln) for ln in lines)
        if len(text.strip()) < 2:
            continue
        text_els.append({"kind": "text", "lines": lines, "text": text,
                         "size": max(sizes) if sizes else 0,
                         "y0": y0, "y1": y1, "x0": x0, "bbox": (x0, y0, x1, y1),
                         "page": page.number, "H": H})
    elems = (text_els
             + [{"kind": "table", "bbox": b, "cells": c, "page": page.number,
                "y0": b[1], "y1": b[3], "x0": b[0]} for b, c in table_data]
             + img_els)
    return order_elements(elems, W)


def is_chrome(text, y0, y1, H, repeat):
    t = text.strip()
    if not t:
        return True
    if _PURE_NUM.match(t):
        return True
    if _RUNHEAD.search(t):
        return True
    if _FOOTNOTE_MARK.match(t):
        return True
    if (y1 < 0.09 * H or y0 > 0.93 * H) and len(t) < 70:
        return True
    if (y1 < 0.12 * H or y0 > 0.90 * H) and len(t) < 95:
        return True
    if repeat >= 3 and len(t) < 60:
        return True
    return False


def is_complete(s):
    s = s.rstrip()
    return bool(s) and s[-1] in _END_PUNCT


_CONT = {"and", "or", "but", "the", "a", "an", "of", "to", "in", "on", "for",
         "with", "that", "which", "who", "this", "these", "those", "by", "as",
         "at", "from", "into", "such", "if", "when", "while", "because",
         "however", "thus", "therefore", "moreover", "furthermore", "also",
         "we", "they", "it", "there", "our", "their", "its", "where", "what",
         "how", "than", "between", "among", "through", "over", "under", "via",
         "using", "based", "due", "e.g", "i.e", "see", "since", "both", "whether"}


def looks_continuation(b):
    b = b.lstrip()
    if not b:
        return False
    if b[0].islower():
        return True
    w = re.match(r"[A-Za-z]+", b)
    return bool(w) and w.group(0).lower() in _CONT


def merge_items(items):
    out = []
    for it in items:
        if it.get("kind") != "text":
            out.append(it)
            continue
        if out and out[-1].get("kind") == "text":
            A, B = out[-1], it
            # never merge headings, list items, or section starts as a unit
            if (A.get("level", 0) in (1, 2, 3)
                    or _LIST_ITEM.match(A["text"])
                    or _LIST_ITEM.match(B["text"])
                    or _SECTION.match(B["text"])):
                out.append(it)
                continue
            # (a) cross-line sentence heal: previous block wasn't a finished
            # sentence and the next continues it
            heal = (not is_complete(A["text"])
                    and (looks_continuation(B["text"])
                         or A["text"].rstrip().endswith((",", ";", "—", "-"))))
            # (b) same-paragraph merge: consecutive body blocks on the SAME
            # page + SAME column with only a normal line gap between them (no
            # blank line) -> join, so a paragraph is not one-sentence-per-line.
            gap = B["y0"] - A["y1"]
            fs = max(A.get("size", 0), B.get("size", 0)) or 11
            same_col = abs(A.get("x0", 0) - B.get("x0", 0)) < 24
            same_page = A.get("page") == B.get("page")
            para = same_page and same_col and -2 <= gap <= fs * 1.6
            if heal or para:
                A["lines"] = (A.get("lines") or []) + (B.get("lines") or [])
                A["text"] = (A["text"].rstrip() + " " + B["text"].strip()).strip()
                A["size"] = max(A["size"], B["size"])
                A["y1"] = max(A["y1"], B["y1"])
                continue
        out.append(it)
    return out


def classify(text, size, median):
    t = text.strip()
    if median and size >= median * 1.5:
        return 1
    if median and size >= median * 1.2:
        return 2
    if _SECTION.match(t):
        return 3
    return 0


def build_items(pdf_path, log):
    doc = fitz.open(pdf_path)
    page_els = []
    for page in doc:
        page_els.append(extract_page(page, page.rect.width, page.rect.height))
    doc.close()

    rep_src = [e["text"].strip() for els in page_els for e in els
               if e["kind"] == "text" and 0 < len(e["text"].strip()) < 60]
    repeat = collections.Counter(rep_src)

    raw, footnotes = [], []
    for els in page_els:
        for e in els:
            if e["kind"] in ("table", "image"):
                raw.append(e)
                continue
            t = e["text"].strip()
            t = _FOOTNOTE_TAIL.sub("", t).strip()
            e["text"] = t
            if is_chrome(t, e["y0"], e["y1"], e["H"], repeat.get(t, 0)):
                for u in _URL.findall(t):
                    if u not in footnotes:
                        footnotes.append(u)
                continue
            urls = _URL.findall(t)
            if urls:
                for u in urls:
                    if u not in footnotes:
                        footnotes.append(u)
                t2 = _URL.sub(" ", t).strip()
                if len(t2) < 15:
                    continue
                e["text"] = t2
            raw.append(e)

    all_sizes = [e["size"] for e in raw if e["kind"] == "text"]
    median = sorted(all_sizes)[len(all_sizes) // 2] if all_sizes else 11
    for e in raw:
        if e["kind"] == "text":
            e["level"] = classify(e["text"], e["size"], median)

    raw = merge_items(raw)

    items, first_text = [], True
    for it in raw:
        if it["kind"] in ("table", "image"):
            items.append(it)
            continue
        text = it["text"].strip()
        if not text:
            continue
        level = it["level"]
        is_title = False
        if first_text and level >= 1 and len(text) < 140:
            is_title = True
            first_text = False
        elif first_text:
            first_text = False
        m = _SECTION_SPLIT.match(text)
        if level in (2, 3) and m and len(text) > 50:
            head = m.group(1).strip()
            body = text[m.end():].strip()
            items.append({"kind": "text", "lines": None, "text": head,
                          "level": 3 if level == 3 else 2, "title": False,
                          "split_head": True, "page": it.get("page", 0)})
            if body:
                items.append({"kind": "text", "lines": None, "text": body,
                              "level": 0, "title": False, "page": it.get("page", 0)})
            continue
        items.append({"kind": "text", "lines": it["lines"], "text": text,
                      "level": level, "title": is_title, "page": it.get("page", 0)})

    nums = []
    for idx, it in enumerate(items):
        if it["kind"] == "text" and it["level"] == 3:
            mm = re.match(r"^(\d+)\.?\s", it["text"])
            if mm:
                nums.append((idx, int(mm.group(1))))
    i = 0
    while i < len(nums):
        j = i
        run_idx = [nums[i][0]]
        while (j + 1 < len(nums)
               and nums[j + 1][1] == nums[j][1] + 1
               and nums[j + 1][0] - nums[j][0] <= 4):
            j += 1
            run_idx.append(nums[j][0])
        if len(run_idx) >= 2:
            for k in run_idx:
                items[k]["level"] = 0
        i = j + 1

    # Merge a detected title with immediately-following short heading lines
    # (multi-line titles like "XXX 建模\n语言研究" should stay one title).
    for i in range(len(items) - 1):
        a, b = items[i], items[i + 1]
        if a is None or b is None:
            continue
        if (a.get("title") and b.get("kind") == "text"
                and b.get("level") in (1, 2, 3)
                and not _SECTION.match(b["text"])
                and 0 < len(b["text"]) < 70):
            a["text"] = (a["text"] + " " + b["text"]).strip()
            a["lines"] = (a.get("lines") or []) + (b.get("lines") or [])
            items[i + 1] = None
    items = [x for x in items if x is not None]

    log.write("  %s : %d items, median=%.1f, footnotes=%d\n" % (
        os.path.basename(pdf_path), len(items), median, len(footnotes)))
    return items, footnotes


# ---------------------------------------------------------------------------
# Translation + document assembly
# ---------------------------------------------------------------------------
def _join_text(zh_lines):
    return " ".join((ln.get("zh") or "") for ln in zh_lines if ln.get("zh")).strip()


def _style_run(r, p):
    if p.style.name.startswith("Heading"):
        sz = 15 if p.style.name == "Heading 1" else 13
        set_run_font(r, HEAD_LATIN, HEAD_CJK, sz)
    else:
        set_run_font(r, BODY_LATIN, BODY_CJK, BODY_SIZE)


def _add_zh_with_math(doc, p, zh, pdf_path, page):
    if not zh:
        return
    r = p.add_run(zh)
    _style_run(r, p)


def _emit_centered_image(doc, pdf_path, page, bbox, zoom=3.0, cap_cm=15, min_cm=0):
    """Render a crop and embed it centered, scaled to its NATURAL on-page width
    but never wider than cap_cm (so wide formulas shrink to fit the page) and
    never narrower than min_cm (keeps tiny formulas legible). Height follows
    width automatically -> no deformation, no oversized blocks."""
    try:
        stream, w_px, h_px = render_crop(pdf_path, page, bbox, zoom)
        nat_w = (w_px / zoom) * 2.54 / 72.0
        if cap_cm and nat_w > cap_cm:
            w = cap_cm
        elif min_cm and nat_w < min_cm:
            w = min_cm
        else:
            w = nat_w
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(stream, width=Cm(w))
        doc.add_paragraph()
    except Exception as e:
        sys.stderr.write("  img/eq render fail: %s\n" % str(e)[:80])


def _emit_text_block(doc, it, heading, title, pdf_path):
    units = it.get("_units")
    if not units:
        units = [{"type": "text", "zh": it.get("text", ""), "mb": []}]
    page = it.get("page", 0)
    if title:
        # title: join all normal-run translations
        zt = " ".join(u.get("zh", "") for u in units if u["type"] == "text")
        p = doc.add_paragraph(style="Title")
        _add_zh_with_math(doc, p, zt, pdf_path, page)
        return

    def newp():
        p = doc.add_paragraph()
        if heading == 1:
            p.style = doc.styles["Heading 1"]
        elif heading == 2:
            p.style = doc.styles["Heading 2"]
        else:
            p.paragraph_format.first_line_indent = Pt(2 * BODY_SIZE)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(6)
        return p

    cur = [None]
    for u in units:
        if u["type"] == "display":
            _emit_centered_image(doc, pdf_path, page, u["bbox"], 3.0, 15, 1.5)
            cur[0] = None
        else:
            if cur[0] is None:
                cur[0] = newp()
            _add_zh_with_math(doc, cur[0], u.get("zh") or "", pdf_path, page)


def _emit_table(doc, it):
    cells = it.get("_cells")
    if not cells:
        return
    nrows = len(cells)
    ncols = max((len(r) for r in cells), default=0)
    if nrows < 1 or ncols < 1:
        return
    try:
        tbl = doc.add_table(rows=nrows, cols=ncols)
        try:
            tbl.style = "Table Grid"
        except Exception:
            pass
        for r, row in enumerate(cells):
            for c in range(ncols):
                val = row[c] if c < len(row) else ""
                cp = tbl.cell(r, c).paragraphs[0]
                run = cp.add_run(val or "")
                set_run_font(run, BODY_LATIN, BODY_CJK, 9)
        doc.add_paragraph()
    except Exception as e:
        sys.stderr.write("  table emit fail: %s\n" % str(e)[:80])


def _fix_heading_dup(doc):
    paras = doc.paragraphs
    for i, p in enumerate(paras):
        if not p.style.name.startswith("Heading"):
            continue
        htxt = p.text.strip()
        aw = htxt.split()
        if not aw:
            continue
        last = aw[-1]
        if i + 1 < len(paras):
            nxt = paras[i + 1]
            if nxt.style.name.startswith("Heading"):
                continue
            for run in nxt.runs:
                if run.text.startswith(last):
                    run.text = run.text[len(last):].lstrip()
                    break


def translate_pdf(pdf_path, out_path, log):
    items, footnotes = build_items(pdf_path, log)

    queue = []
    meta = []
    for ii, it in enumerate(items):
        if it["kind"] != "text" or it.get("split_head"):
            continue
        lines = it["lines"] or []
        units = []
        buf = {"type": "text", "text": "", "mb": []}
        for line in lines:
            if line_is_display(line):
                if buf["text"]:
                    units.append(buf)
                    buf = {"type": "text", "text": "", "mb": []}
                mspans = [s for s in line if span_math(s)]
                bb = union_bbox(mspans) if mspans else union_bbox(line)
                units.append({"type": "display", "bbox": bb})
            else:
                s, mb = prep_line(line)
                if s:
                    buf["text"] = (buf["text"] + " " + s).strip()
                    buf["mb"] = buf["mb"] + mb
        if buf["text"]:
            units.append(buf)
        it["_units"] = units
        for ui, u in enumerate(units):
            if u["type"] == "text" and u["text"]:
                queue.append(u["text"])
                meta.append((ii, ui))

    cell_meta = []
    for ii, it in enumerate(items):
        if it["kind"] != "table" or not it.get("cells"):
            continue
        cells = it["cells"]
        ncols = max((len(row) for row in cells), default=0)
        it["_cells"] = [[(cells[r][c] or "").strip() if c < len(cells[r]) else ""
                         for c in range(ncols)]
                        for r, row in enumerate(cells)]
        for r, row in enumerate(it["_cells"]):
            for c, val in enumerate(row):
                if val:
                    queue.append(val)
                    cell_meta.append((ii, r, c))

    with ThreadPoolExecutor(max_workers=12) as ex:
        results = list(ex.map(gtranslate, queue))

    q = 0
    for ii, ui in meta:
        items[ii]["_units"][ui]["zh"] = clean(results[q])
        q += 1
    for ii, r, c in cell_meta:
        items[ii]["_cells"][r][c] = clean(results[q])
        q += 1

    doc = Document()
    set_cjk(doc.styles["Normal"], BODY_LATIN, BODY_CJK, BODY_SIZE)
    set_cjk(doc.styles["Heading 1"], HEAD_LATIN, HEAD_CJK, 15)
    set_cjk(doc.styles["Heading 2"], HEAD_LATIN, HEAD_CJK, 13)
    set_cjk(doc.styles["Title"], HEAD_LATIN, HEAD_CJK, 18)

    for it in items:
        if it["kind"] == "image":
            _emit_centered_image(doc, pdf_path, it["page"], it["bbox"], 2.0, 15, 6)
            continue
        if it["kind"] == "table":
            if it.get("cells"):
                _emit_table(doc, it)
            else:
                _emit_centered_image(doc, pdf_path, it["page"], it["bbox"], 2.2, 14, 0)
            continue
        if it.get("split_head"):
            zh = it["text"]
            if it["level"] in (1, 3) and not re.search(r"\.\d", it["text"]):
                doc.add_heading(zh, level=1)
            elif it["level"] in (2, 3):
                doc.add_heading(zh, level=2)
            continue
        title = it.get("title")
        if it["level"] in (1, 3) and not re.search(r"\.\d", it["text"]):
            _emit_text_block(doc, it, 1, title, pdf_path)
        elif it["level"] in (2, 3):
            _emit_text_block(doc, it, 2, title, pdf_path)
        else:
            _emit_text_block(doc, it, 0, title, pdf_path)

    _fix_heading_dup(doc)

    if footnotes:
        doc.add_heading("参考链接", level=2)
        for u in footnotes:
            fp = doc.add_paragraph()
            fr = fp.add_run(u)
            fr.font.size = Pt(9)
            fr.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
    doc.save(out_path)
    if pdf_path in _PDF_CACHE:
        try:
            _PDF_CACHE[pdf_path].close()
        except Exception:
            pass
        del _PDF_CACHE[pdf_path]


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--force", action="store_true")
    ap.add_argument("paths", nargs="*", default=["."])
    args = ap.parse_args()
    paths = args.paths or ["."]
    pdfs = []
    for p in paths:
        if os.path.isdir(p):
            for f in sorted(os.listdir(p)):
                if f.lower().endswith(".pdf") and "-翻译" not in f:
                    pdfs.append(os.path.join(p, f))
        elif p.lower().endswith(".pdf"):
            pdfs.append(p)
    print("Total PDFs to translate: %d" % len(pdfs))
    log = open("_translate_batch.log", "a", encoding="utf-8")
    for i, pdf in enumerate(pdfs, 1):
        out = os.path.splitext(pdf)[0] + "-翻译" + OUT_EXT
        if os.path.exists(out) and not args.force:
            print("[%d/%d] SKIP (exists): %s" % (i, len(pdfs),
                                                 os.path.basename(pdf)))
            continue
        t0 = time.time()
        try:
            translate_pdf(pdf, out, log)
            log.flush()
            print("[%d/%d] OK  %.1fs  %s" % (i, len(pdfs), time.time() - t0,
                                             os.path.basename(pdf)))
        except Exception as e:
            print("[%d/%d] FAIL %s : %s" % (i, len(pdfs),
                                            os.path.basename(pdf), e))
    log.close()


if __name__ == "__main__":
    main()
